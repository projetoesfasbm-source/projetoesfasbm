# backend/services/mail_merge_service.py

import pandas as pd
from docx import Document
from io import BytesIO
import zipfile
import re

class MailMergeService:
    @staticmethod
    def _replace_text_in_paragraph(paragraph, replacements):
        """
        Substitui múltiplos placeholders em um parágrafo, preservando a formatação dominante.
        Esta função foi corrigida para ser mais robusta e eficiente.
        """
        # 1. Constrói o texto completo do parágrafo
        full_text = "".join(run.text for run in paragraph.runs)
        
        # 2. Verifica se há algo para substituir antes de fazer qualquer alteração
        if not any(placeholder in full_text for placeholder in replacements.keys()):
            return

        # 3. Salva a formatação do primeiro 'run' (que geralmente define o estilo da linha)
        # e a formatação do parágrafo (alinhamento)
        original_run = paragraph.runs[0] if paragraph.runs else None
        p_format = paragraph.paragraph_format

        # 4. Executa TODAS as substituições na string de texto
        modified_text = full_text
        for placeholder, value in replacements.items():
            modified_text = modified_text.replace(placeholder, str(value))

        # 5. Limpa os 'runs' antigos do parágrafo
        for run in paragraph.runs:
            p = run._element
            p.getparent().remove(p)

        # 6. Adiciona o texto modificado de volta em um novo 'run'
        new_run = paragraph.add_run(modified_text)
        
        # 7. Reaplica a formatação salva ao novo 'run' e ao parágrafo
        if original_run:
            new_run.style = original_run.style
            font = original_run.font
            new_run.font.name = font.name
            new_run.font.size = font.size
            new_run.font.bold = font.bold
            new_run.font.italic = font.italic
            new_run.font.underline = font.underline
            new_run.font.color.rgb = font.color.rgb
        
        paragraph.paragraph_format.alignment = p_format.alignment

    @staticmethod
    def generate_documents(template_file, data_file, output_format='docx'):
        """
        Gera documentos individuais, preservando formatação e orientação de página,
        e os compacta em um arquivo ZIP.
        """
        try:
            template_buffer = BytesIO(template_file.read())
            df = pd.read_excel(data_file)
            records = df.astype(str).to_dict('records')

            zip_buffer = BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
                for index, record in enumerate(records):
                    # Recarrega o template do buffer para cada novo documento.
                    # Isso é crucial para preservar a estrutura original, incluindo orientações de página.
                    template_buffer.seek(0)
                    doc = Document(template_buffer)

                    # Cria o dicionário de substituições para o registro atual
                    replacements = {f"{{{{{key.strip()}}}}}": str(value) for key, value in record.items()}

                    # Itera e substitui nos parágrafos e tabelas
                    for p in doc.paragraphs:
                        MailMergeService._replace_text_in_paragraph(p, replacements)
                    
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    MailMergeService._replace_text_in_paragraph(p, replacements)

                    # Salva o documento individual em memória
                    doc_buffer = BytesIO()
                    doc.save(doc_buffer)
                    doc_buffer.seek(0)
                    
                    file_name_base = record.get('nome', f"documento_{index + 1}").strip().replace(" ", "_")
                    filename = f"Certificado_{file_name_base}.docx"
                    
                    zf.writestr(filename, doc_buffer.getvalue())

            zip_buffer.seek(0)
            return zip_buffer, None
            
        except Exception as e:
            error_message = f"Ocorreu um erro: {str(e)}. Verifique se os placeholders no seu template (ex: {{nome}}) correspondem às colunas da planilha."
            return None, error_message